"""Generate a Word (.docx) file from a podcast markdown table.

Usage:
    python build_docx.py <input.md> [output.docx]

If no output path is given, the .docx is saved next to the input file.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(h)
    for run in h.runs:
        run.font.name = "David"
    return h


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "David"
    return p


def style_cell(cell, text, bold=False, size=11):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "David"


def parse_md_table(md_path: Path):
    """Parse a markdown file and extract header metadata and table rows.

    Returns (meta: dict, rows: list of (num, speaker, content, note)).
    """
    text = md_path.read_text(encoding="utf-8")

    # Extract frontmatter-style metadata from lines starting with **Key:**
    meta = {}
    for line in text.splitlines():
        m = re.match(r"\*\*(.+?):\*\*\s*(.+)", line)
        if m:
            meta[m.group(1).strip()] = m.group(2).strip()

    # Extract title from first # heading
    title_match = re.search(r"^#\s+(.+)", text, re.MULTILINE)
    if title_match:
        meta.setdefault("title", title_match.group(1).strip())

    # Parse markdown table rows (skip header and separator rows)
    rows = []
    for line in text.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[-| ]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if len(cells) < 4:
            continue
        # Skip header row (non-numeric first cell that is not empty)
        if cells[0] and not cells[0].isdigit():
            continue
        num = cells[0]
        speaker = cells[1]
        content = cells[2]
        note = cells[3] if len(cells) > 3 else ""
        rows.append((num, speaker, content, note))

    return meta, rows


def build_doc(md_path: Path, out_path: Path):
    meta, rows = parse_md_table(md_path)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "David"
    style.font.size = Pt(11)

    title = meta.get("title", md_path.stem)
    add_heading(doc, title, level=0)

    if "קורס" in meta:
        add_para(doc, f"קורס: {meta['קורס']}", bold=True)
    if "מבוסס על" in meta:
        add_para(doc, f"מבוסס על: {meta['מבוסס על']}")
    if "זמן האזנה משוער" in meta:
        add_para(doc, f"זמן האזנה משוער: {meta['זמן האזנה משוער']}")
    if "שחקנים" in meta:
        add_para(doc, f"שחקנים: {meta['שחקנים']}")

    add_para(
        doc,
        'הערה: במקומות המסומנים [ציטוט תורני – ליועצות הדת] ניתן להשלים מקור מתאים.',
    )

    add_heading(doc, "תמלול השיחה המלא", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"

    tbl = table._tbl
    tblPr = tbl.tblPr
    bidiVisual = OxmlElement("w:bidiVisual")
    tblPr.append(bidiVisual)

    headers = ["#", "דובר", "תוכן (מה שומעים)", "הערות הפקה"]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        style_cell(hdr[i], h, bold=True, size=11)

    for num, speaker, content, note in rows:
        row_cells = table.add_row().cells
        style_cell(row_cells[0], num)
        style_cell(row_cells[1], speaker, bold=bool(speaker))
        style_cell(row_cells[2], content)
        style_cell(row_cells[3], note)

    doc.save(out_path)
    print(f"Saved: {out_path} ({len(rows)} rows)")


def main():
    if len(sys.argv) < 2:
        # Backward-compatible: build episode 1 from its .md file
        default_md = Path(__file__).parent / "podcast-01-takhilato-shel-masa.md"
        if default_md.exists():
            out = default_md.with_suffix(".docx")
            build_doc(default_md, out)
        else:
            print("Usage: python build_docx.py <input.md> [output.docx]")
            sys.exit(1)
        return

    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"Error: file not found: {md_path}")
        sys.exit(1)

    out_path = Path(sys.argv[2]) if len(sys.argv) > 2 else md_path.with_suffix(".docx")
    build_doc(md_path, out_path)


if __name__ == "__main__":
    main()
