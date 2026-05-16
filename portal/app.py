"""L-net Podcast Studio — local web portal."""

import io
import json
import re
import sqlite3
import sys
from pathlib import Path

from flask import (Flask, Response, jsonify, redirect, render_template,
                   request, send_file, url_for)

BASE_DIR = Path(__file__).parent
DB_PATH = BASE_DIR / "podcast_studio.db"
OUTPUT_DIR = BASE_DIR.parent / "output"

WORDS_PER_MINUTE = 120

app = Flask(__name__)


# ---------------------------------------------------------------------------
# Script analysis helpers
# ---------------------------------------------------------------------------

def count_script_words(md_table: str) -> int:
    """Sum word count of the 'תוכן' column (column 3) across all rows."""
    if not md_table:
        return 0
    total = 0
    for line in md_table.split("\n"):
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[-| ]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if len(cells) < 3:
            continue
        # Skip header row (first cell is non-numeric and non-empty)
        if cells[0] and not cells[0].isdigit():
            continue
        total += len(cells[2].split())
    return total


def estimate_duration(word_count: int) -> str:
    minutes = max(1, round(word_count / WORDS_PER_MINUTE))
    return f"כ-{minutes} דקות"


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file (paragraphs + table cells)."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


@app.template_filter("safe_nl")
def safe_nl(text: str) -> str:
    """Convert newlines to <br> for display."""
    import markupsafe
    return markupsafe.Markup(
        markupsafe.escape(text or "").replace("\n", markupsafe.Markup("<br>"))
    )


@app.template_filter("render_table")
def render_table(md: str) -> str:
    """Render a markdown table as an HTML table, or fallback to pre."""
    import markupsafe
    if not md:
        return ""
    lines = [l for l in md.split("\n") if l.strip().startswith("|")]
    if len(lines) < 2:
        return markupsafe.Markup(
            f'<pre style="direction:rtl;white-space:pre-wrap">{markupsafe.escape(md)}</pre>'
        )
    data_lines = [l for l in lines if not l.strip().replace("|", "").replace("-", "").replace(" ", "")]
    table_lines = [l for l in lines if l not in data_lines]

    def cells(line):
        return [c.strip() for c in line.strip().strip("|").split("|")]

    html = '<div class="script-table-wrap"><table class="script-table"><thead><tr>'
    for h in cells(table_lines[0]):
        html += f"<th>{markupsafe.escape(h)}</th>"
    html += "</tr></thead><tbody>"
    for row in table_lines[1:]:
        cs = cells(row)
        html += "<tr>"
        for i, c in enumerate(cs):
            cls = ' class="speaker"' if i == 1 and c else ""
            html += f"<td{cls}>{markupsafe.escape(c)}</td>"
        html += "</tr>"
    html += "</tbody></table></div>"
    return markupsafe.Markup(html)


# ---------------------------------------------------------------------------
# DB
# ---------------------------------------------------------------------------

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def init_db():
    with get_db() as db:
        db.executescript("""
            CREATE TABLE IF NOT EXISTS cluster (
                id   INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL UNIQUE
            );
            CREATE TABLE IF NOT EXISTS course (
                id         INTEGER PRIMARY KEY AUTOINCREMENT,
                cluster_id INTEGER NOT NULL REFERENCES cluster(id) ON DELETE CASCADE,
                name       TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS episode (
                id        INTEGER PRIMARY KEY AUTOINCREMENT,
                course_id INTEGER NOT NULL REFERENCES course(id) ON DELETE CASCADE,
                title     TEXT NOT NULL,
                status    TEXT NOT NULL DEFAULT 'pending'
            );
            CREATE TABLE IF NOT EXISTS workspace (
                episode_id   INTEGER PRIMARY KEY REFERENCES episode(id) ON DELETE CASCADE,
                course_name  TEXT DEFAULT '',
                raw_content  TEXT DEFAULT '',
                parashiyot   TEXT DEFAULT '',
                avoid_notes  TEXT DEFAULT '',
                analysis     TEXT DEFAULT '',
                draft        TEXT DEFAULT '',
                critique     TEXT DEFAULT '',
                revised      TEXT DEFAULT '',
                step         INTEGER DEFAULT 0,
                export_path  TEXT DEFAULT ''
            );
        """)


def get_or_create_workspace(db, episode_id: int, course_name: str = ""):
    row = db.execute(
        "SELECT * FROM workspace WHERE episode_id = ?", (episode_id,)
    ).fetchone()
    if not row:
        db.execute(
            "INSERT INTO workspace (episode_id, course_name) VALUES (?, ?)",
            (episode_id, course_name),
        )
        db.commit()
        row = db.execute(
            "SELECT * FROM workspace WHERE episode_id = ?", (episode_id,)
        ).fetchone()
    return row


# ---------------------------------------------------------------------------
# Dashboard
# ---------------------------------------------------------------------------

@app.get("/")
def index():
    db = get_db()
    clusters = db.execute("SELECT * FROM cluster ORDER BY id").fetchall()
    data = []
    for cl in clusters:
        courses = db.execute(
            "SELECT * FROM course WHERE cluster_id = ? ORDER BY id", (cl["id"],)
        ).fetchall()
        course_data = []
        total_eps = written_eps = 0
        for co in courses:
            episodes = db.execute(
                "SELECT * FROM episode WHERE course_id = ? ORDER BY id", (co["id"],)
            ).fetchall()
            total_eps += len(episodes)
            written_eps += sum(1 for e in episodes if e["status"] == "written")
            course_data.append({"course": co, "episodes": episodes})
        data.append({
            "cluster": cl,
            "courses": course_data,
            "total": total_eps,
            "written": written_eps,
        })
    return render_template("index.html", data=data)


# ---------------------------------------------------------------------------
# Clusters
# ---------------------------------------------------------------------------

@app.post("/cluster/add")
def cluster_add():
    name = request.form.get("name", "").strip()
    if name:
        try:
            with get_db() as db:
                db.execute("INSERT INTO cluster (name) VALUES (?)", (name,))
        except sqlite3.IntegrityError:
            pass
    return redirect(url_for("index"))


@app.post("/cluster/<int:cluster_id>/delete")
def cluster_delete(cluster_id):
    with get_db() as db:
        db.execute("DELETE FROM cluster WHERE id = ?", (cluster_id,))
    return redirect(url_for("index"))


# ---------------------------------------------------------------------------
# Courses
# ---------------------------------------------------------------------------

@app.post("/cluster/<int:cluster_id>/course/add")
def course_add(cluster_id):
    name = request.form.get("name", "").strip()
    if name:
        with get_db() as db:
            db.execute(
                "INSERT INTO course (cluster_id, name) VALUES (?, ?)",
                (cluster_id, name),
            )
    return redirect(url_for("index") + f"#cluster-{cluster_id}")


@app.post("/course/<int:course_id>/delete")
def course_delete(course_id):
    db = get_db()
    row = db.execute(
        "SELECT cluster_id FROM course WHERE id = ?", (course_id,)
    ).fetchone()
    cluster_id = row["cluster_id"] if row else None
    with db:
        db.execute("DELETE FROM course WHERE id = ?", (course_id,))
    return redirect(url_for("index") + (f"#cluster-{cluster_id}" if cluster_id else ""))


# ---------------------------------------------------------------------------
# Episodes
# ---------------------------------------------------------------------------

@app.post("/course/<int:course_id>/episode/add")
def episode_add(course_id):
    title = request.form.get("title", "").strip()
    if title:
        with get_db() as db:
            db.execute(
                "INSERT INTO episode (course_id, title) VALUES (?, ?)",
                (course_id, title),
            )
    db = get_db()
    row = db.execute(
        "SELECT cluster_id FROM course WHERE id = ?", (course_id,)
    ).fetchone()
    cluster_id = row["cluster_id"] if row else None
    return redirect(url_for("index") + (f"#cluster-{cluster_id}" if cluster_id else ""))


@app.post("/episode/<int:episode_id>/toggle")
def episode_toggle(episode_id):
    db = get_db()
    row = db.execute(
        "SELECT status, course_id FROM episode WHERE id = ?", (episode_id,)
    ).fetchone()
    if not row:
        return redirect(url_for("index"))
    new_status = "written" if row["status"] == "pending" else "pending"
    with db:
        db.execute(
            "UPDATE episode SET status = ? WHERE id = ?", (new_status, episode_id)
        )
    row2 = db.execute(
        "SELECT cluster_id FROM course WHERE id = ?", (row["course_id"],)
    ).fetchone()
    cluster_id = row2["cluster_id"] if row2 else None
    return redirect(url_for("index") + (f"#cluster-{cluster_id}" if cluster_id else ""))


@app.post("/episode/<int:episode_id>/delete")
def episode_delete(episode_id):
    db = get_db()
    row = db.execute(
        "SELECT c.cluster_id FROM episode e JOIN course c ON e.course_id = c.id WHERE e.id = ?",
        (episode_id,),
    ).fetchone()
    cluster_id = row["cluster_id"] if row else None
    with db:
        db.execute("DELETE FROM episode WHERE id = ?", (episode_id,))
    return redirect(url_for("index") + (f"#cluster-{cluster_id}" if cluster_id else ""))


# ---------------------------------------------------------------------------
# Workspace — main page
# ---------------------------------------------------------------------------

@app.get("/episode/<int:episode_id>/workspace")
def workspace(episode_id):
    db = get_db()
    ep = db.execute("SELECT * FROM episode WHERE id = ?", (episode_id,)).fetchone()
    if not ep:
        return redirect(url_for("index"))
    co = db.execute("SELECT * FROM course WHERE id = ?", (ep["course_id"],)).fetchone()
    cl = db.execute(
        "SELECT * FROM cluster WHERE id = ?", (co["cluster_id"],)
    ).fetchone()
    ws = get_or_create_workspace(db, episode_id, co["name"] if co else "")

    # Compute word counts + durations for any populated draft/revised
    draft_wc = count_script_words(ws["draft"]) if ws["draft"] else 0
    revised_table = (ws["revised"] or "").split("---")[0] if ws["revised"] else ""
    revised_wc = count_script_words(revised_table) if revised_table else 0

    return render_template(
        "workspace.html",
        episode=ep,
        course=co,
        cluster=cl,
        ws=ws,
        draft_wc=draft_wc,
        draft_duration=estimate_duration(draft_wc) if draft_wc else "",
        revised_wc=revised_wc,
        revised_duration=estimate_duration(revised_wc) if revised_wc else "",
        words_per_minute=WORDS_PER_MINUTE,
    )


# ---------------------------------------------------------------------------
# Workspace API — each step calls Claude
# ---------------------------------------------------------------------------

def _ws_error(msg: str):
    return jsonify({"ok": False, "error": msg}), 500


@app.post("/episode/<int:episode_id>/api/upload-content")
def api_upload_content(episode_id):
    """Receive a .docx file, extract text, return it (and save to workspace)."""
    if "file" not in request.files:
        return _ws_error("לא נשלח קובץ")
    file = request.files["file"]
    if not file.filename:
        return _ws_error("שם קובץ ריק")
    if not file.filename.lower().endswith(".docx"):
        return _ws_error("רק קבצי .docx נתמכים")

    try:
        text = extract_docx_text(file.read())
    except Exception as e:
        return _ws_error(f"שגיאה בקריאת הקובץ: {e}")

    if not text.strip():
        return _ws_error("הקובץ ריק או שאין בו טקסט קריא")

    # Save extracted text to the workspace
    with get_db() as db:
        db.execute(
            """INSERT INTO workspace (episode_id, raw_content) VALUES (?, ?)
               ON CONFLICT(episode_id) DO UPDATE SET raw_content=excluded.raw_content""",
            (episode_id, text),
        )

    return jsonify({
        "ok": True,
        "text": text,
        "filename": file.filename,
        "char_count": len(text),
    })


@app.post("/episode/<int:episode_id>/api/analyze")
def api_analyze(episode_id):
    try:
        from claude_client import analyze_content
    except ImportError as e:
        return _ws_error(str(e))

    data = request.get_json()
    raw = data.get("raw_content", "").strip()
    if not raw:
        return _ws_error("יש להזין חומר גולמי")

    parashiyot = data.get("parashiyot", "")
    avoid = data.get("avoid_notes", "")

    try:
        result = analyze_content(raw, parashiyot, avoid)
    except Exception as e:
        return _ws_error(str(e))

    with get_db() as db:
        db.execute(
            """INSERT INTO workspace (episode_id, raw_content, parashiyot, avoid_notes, analysis, step)
               VALUES (?, ?, ?, ?, ?, 2)
               ON CONFLICT(episode_id) DO UPDATE SET
                 raw_content=excluded.raw_content,
                 parashiyot=excluded.parashiyot,
                 avoid_notes=excluded.avoid_notes,
                 analysis=excluded.analysis,
                 step=MAX(step, 2)""",
            (episode_id, raw, parashiyot, avoid, result),
        )
    return jsonify({"ok": True, "result": result})


@app.post("/episode/<int:episode_id>/api/draft")
def api_draft(episode_id):
    try:
        from claude_client import write_draft
    except ImportError as e:
        return _ws_error(str(e))

    db = get_db()
    ws = db.execute(
        "SELECT * FROM workspace WHERE episode_id = ?", (episode_id,)
    ).fetchone()
    ep = db.execute("SELECT * FROM episode WHERE id = ?", (episode_id,)).fetchone()
    co = db.execute(
        "SELECT * FROM course WHERE id = ?", (ep["course_id"],)
    ).fetchone() if ep else None

    if not ws or not ws["analysis"]:
        return _ws_error("יש לבצע ניתוח תרבותי קודם")

    try:
        result = write_draft(
            ws["raw_content"],
            ws["analysis"],
            co["name"] if co else "",
            ep["title"] if ep else "",
        )
    except Exception as e:
        return _ws_error(str(e))

    with get_db() as db:
        db.execute(
            "UPDATE workspace SET draft=?, step=MAX(step,3) WHERE episode_id=?",
            (result, episode_id),
        )
    wc = count_script_words(result)
    return jsonify({
        "ok": True,
        "result": result,
        "word_count": wc,
        "duration": estimate_duration(wc),
    })


@app.post("/episode/<int:episode_id>/api/critique")
def api_critique(episode_id):
    try:
        from claude_client import critique_draft
    except ImportError as e:
        return _ws_error(str(e))

    db = get_db()
    ws = db.execute(
        "SELECT * FROM workspace WHERE episode_id = ?", (episode_id,)
    ).fetchone()
    if not ws or not ws["draft"]:
        return _ws_error("יש לכתוב טיוטה קודם")

    data = request.get_json() or {}
    draft_to_critique = data.get("draft", ws["draft"])

    try:
        result = critique_draft(draft_to_critique)
    except Exception as e:
        return _ws_error(str(e))

    with get_db() as db:
        db.execute(
            "UPDATE workspace SET draft=?, critique=?, step=MAX(step,4) WHERE episode_id=?",
            (draft_to_critique, result, episode_id),
        )
    return jsonify({"ok": True, "result": result})


@app.post("/episode/<int:episode_id>/api/revise")
def api_revise(episode_id):
    try:
        from claude_client import revise_draft
    except ImportError as e:
        return _ws_error(str(e))

    db = get_db()
    ws = db.execute(
        "SELECT * FROM workspace WHERE episode_id = ?", (episode_id,)
    ).fetchone()
    if not ws or not ws["critique"]:
        return _ws_error("יש לקבל ביקורת קודם")

    try:
        result = revise_draft(ws["draft"], ws["critique"])
    except Exception as e:
        return _ws_error(str(e))

    with get_db() as db:
        db.execute(
            "UPDATE workspace SET revised=?, step=MAX(step,5) WHERE episode_id=?",
            (result, episode_id),
        )
    table_part = result.split("---")[0] if "---" in result else result
    wc = count_script_words(table_part)
    return jsonify({
        "ok": True,
        "result": result,
        "word_count": wc,
        "duration": estimate_duration(wc),
    })


@app.post("/episode/<int:episode_id>/api/export")
def api_export(episode_id):
    db = get_db()
    ws = db.execute(
        "SELECT * FROM workspace WHERE episode_id = ?", (episode_id,)
    ).fetchone()
    ep = db.execute("SELECT * FROM episode WHERE id = ?", (episode_id,)).fetchone()
    if not ws or not (ws["revised"] or ws["draft"]):
        return _ws_error("אין תסריט מוכן לייצוא")

    script = ws["revised"] or ws["draft"]

    # Extract table section only (in case there's a change report after ---)
    if "---" in script:
        script = script.split("---")[0].strip()

    # Build a safe filename from episode title
    safe_title = (ep["title"] if ep else f"episode-{episode_id}")
    safe_title = (
        safe_title.replace(" ", "-")
        .replace("/", "-")
        .replace("\\", "-")
        .replace("—", "")
        .replace(":", "")
    )[:60]
    md_filename = f"podcast-{episode_id:02d}-{safe_title}.md"
    docx_filename = md_filename.replace(".md", ".docx")
    md_path = OUTPUT_DIR / md_filename
    docx_path = OUTPUT_DIR / docx_filename

    # Write metadata header + script to .md
    co = db.execute(
        "SELECT * FROM course WHERE id = ?", (ep["course_id"],)
    ).fetchone() if ep else None
    word_count = count_script_words(script)
    duration = estimate_duration(word_count)

    header = f"# פודקאסט {episode_id}: {ep['title'] if ep else ''}\n\n"
    if co:
        header += f"**קורס:** {co['name']}\n"
    header += f"**זמן האזנה משוער:** {duration}\n"
    header += f"**מספר מילים:** {word_count} (לפי {WORDS_PER_MINUTE} מילים לדקה)\n"
    header += f"**שחקנים:** יצחק לוי / חיים מנחם\n\n---\n\n"
    md_path.write_text(header + script, encoding="utf-8")

    # Generate .docx
    try:
        sys.path.insert(0, str(BASE_DIR.parent / "output"))
        from build_docx import build_doc
        build_doc(md_path, docx_path)
    except Exception as e:
        return _ws_error(f"שגיאה בייצוא: {e}")

    # Mark episode as written
    with get_db() as db:
        db.execute(
            "UPDATE workspace SET export_path=?, step=MAX(step,6) WHERE episode_id=?",
            (str(docx_path), episode_id),
        )
        db.execute(
            "UPDATE episode SET status='written' WHERE id=?", (episode_id,)
        )

    return jsonify({
        "ok": True,
        "md_path": str(md_path),
        "docx_filename": docx_filename,
        "download_url": url_for("download_docx", episode_id=episode_id),
    })


@app.get("/episode/<int:episode_id>/download")
def download_docx(episode_id):
    db = get_db()
    ws = db.execute(
        "SELECT export_path FROM workspace WHERE episode_id = ?", (episode_id,)
    ).fetchone()
    if not ws or not ws["export_path"]:
        return "קובץ לא נמצא", 404
    path = Path(ws["export_path"])
    if not path.exists():
        return "קובץ לא נמצא", 404
    return send_file(path, as_attachment=True, download_name=path.name)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    init_db()
    print("\n  L-net Podcast Studio")
    print("  http://localhost:5050\n")
    app.run(debug=True, port=5050, threaded=True)
